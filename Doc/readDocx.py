# readDocx.py
# Package python-docx needs to be installed

import docx

#file = 'doc/Mail Template.docx'

#print(len(doc.paragraphs))

def getText(fileName):
    doc = docx.Document(fileName)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para)
    return fullText

#print(getText(file))

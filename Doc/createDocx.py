# createDocx.py
# This file creates a word doc file
# Package python-docx needs to be installed

import docx

doc = docx.Document()

doc.add_paragraph('Hello Ujjwal','Title')
doc.add_paragraph('Heading','Heading 1')
doc.add_paragraph('Body','Normal')

doc.save('doc/test.docx')